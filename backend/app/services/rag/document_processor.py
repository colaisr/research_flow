"""
Document processing service for RAG.
Handles text extraction from various file types and chunking.
"""
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import logging
import re

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing documents (text extraction, chunking)."""
    
    def __init__(self, chunk_size_tokens: int = 800, chunk_overlap_tokens: int = 150):
        """Initialize document processor.
        
        Args:
            chunk_size_tokens: Target chunk size in tokens (default: 800)
            chunk_overlap_tokens: Overlap between chunks in tokens (default: 150)
        """
        self.chunk_size_tokens = chunk_size_tokens
        self.chunk_overlap_tokens = chunk_overlap_tokens
    
    def extract_text_from_file(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a file based on its extension.
        
        PDFs and images use AI OCR via OpenRouter (GPT-4o Vision).
        No local dependencies required (no Tesseract/Poppler needed).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        metadata = {
            "filename": file_path.name,
            "file_size": file_path.stat().st_size,
            "file_type": file_ext,
            "ocr_method": "ai",  # Always AI via OpenRouter
        }
        
        try:
            if file_ext == ".pdf":
                # PDFs always use AI OCR via OpenRouter
                text, pdf_metadata = self._extract_text_from_pdf(file_path)
                metadata.update(pdf_metadata)
            elif file_ext in [".docx", ".doc"]:
                text = self._extract_text_from_docx(file_path)
            elif file_ext == ".txt":
                text = self._extract_text_from_txt(file_path)
            elif file_ext in [".html", ".htm"]:
                text = self._extract_text_from_html(file_path)
            elif file_ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"]:
                # Image files - use AI OCR via OpenRouter
                text, image_metadata = self._extract_text_from_image(file_path)
                metadata.update(image_metadata)
            elif file_ext in [".xlsx", ".xls"]:
                # Excel files - extract all sheets as formatted tables
                text, excel_metadata = self._extract_text_from_excel(file_path)
                metadata.update(excel_metadata)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            metadata["text_length"] = len(text)
            metadata["extraction_success"] = True
            
            return text, metadata
        
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            metadata["extraction_success"] = False
            metadata["error"] = str(e)
            raise
    
    def extract_text_from_url(self, url: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a URL (HTML content).
        
        Args:
            url: URL to fetch
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        import httpx
        from bs4 import BeautifulSoup
        
        metadata = {
            "url": url,
            "source_type": "url",
        }
        
        try:
            # Fetch URL content
            with httpx.Client(timeout=30.0) as client:
                response = client.get(url, follow_redirects=True)
                response.raise_for_status()
                
                # Parse HTML
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style", "nav", "footer", "header"]):
                    script.decompose()
                
                # Extract text
                text = soup.get_text(separator='\n', strip=True)
                
                # Clean up multiple newlines
                text = re.sub(r'\n\s*\n', '\n\n', text)
                
                metadata["text_length"] = len(text)
                metadata["extraction_success"] = True
                metadata["status_code"] = response.status_code
                
                return text, metadata
        
        except Exception as e:
            logger.error(f"Failed to extract text from URL {url}: {e}")
            metadata["extraction_success"] = False
            metadata["error"] = str(e)
            raise
    
    def _extract_text_from_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file with tables and structure preservation.
        
        Note: PDFs always use AI OCR via OpenRouter.
        This provides better accuracy and no local dependencies.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        import pdfplumber
        
        text_parts = []
        metadata = {
            "page_count": 0,
            "table_count": 0,
            "has_tables": False,
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)
                
                # Extract PDF metadata if available
                if pdf.metadata:
                    if "Title" in pdf.metadata:
                        metadata["pdf_title"] = pdf.metadata["Title"]
                    if "Author" in pdf.metadata:
                        metadata["pdf_author"] = pdf.metadata["Author"]
                    if "Subject" in pdf.metadata:
                        metadata["pdf_subject"] = pdf.metadata["Subject"]
                    if "Creator" in pdf.metadata:
                        metadata["pdf_creator"] = pdf.metadata["Creator"]
                    if "Producer" in pdf.metadata:
                        metadata["pdf_producer"] = pdf.metadata["Producer"]
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_parts = []
                    
                    # Extract text with structure preservation
                    page_text = self._extract_text_with_structure(page)
                    
                    # Check if page is scanned (no or very little text)
                    is_scanned = not page_text or len(page_text.strip()) < 50
                    
                    if is_scanned:
                        # Use AI OCR for scanned pages (always use AI for PDFs)
                        logger.info(f"Page {page_num} appears to be scanned, using AI OCR via OpenRouter")
                        try:
                            ocr_text = self._extract_text_with_vision_api(file_path, page_num)
                            
                            if ocr_text:
                                page_parts.append(ocr_text)
                                metadata["ocr_used"] = True
                                metadata["ocr_method"] = "ai"  # Always AI for PDFs
                                metadata.setdefault("ocr_pages", []).append(page_num)
                        except Exception as ocr_error:
                            logger.warning(f"AI OCR failed for page {page_num}: {ocr_error}")
                            # Continue with empty page if OCR fails
                    elif page_text:
                        page_parts.append(page_text)
                    
                    # Extract tables from the page (only if not scanned)
                    if not is_scanned:
                        try:
                            tables = page.extract_tables()
                            if tables:
                                metadata["has_tables"] = True
                                for table_num, table in enumerate(tables, 1):
                                    try:
                                        metadata["table_count"] += 1
                                        table_text = self._format_table(table)
                                        if table_text:
                                            page_parts.append(f"\n[Таблица {table_num} на странице {page_num}]\n{table_text}")
                                    except Exception as table_error:
                                        logger.warning(f"Failed to format table {table_num} on page {page_num}: {table_error}")
                                        # Continue with other tables
                                        continue
                        except Exception as table_extract_error:
                            logger.warning(f"Failed to extract tables from page {page_num}: {table_extract_error}")
                            # Continue processing text even if table extraction fails
                    
                    if page_parts:
                        page_content = "\n\n".join(page_parts)
                        text_parts.append(f"--- Страница {page_num} ---\n{page_content}")
        
        except FileNotFoundError:
            error_msg = f"PDF file not found: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
        except PermissionError:
            error_msg = f"Permission denied when reading PDF file: {file_path}"
            logger.error(error_msg)
            raise PermissionError(error_msg)
        except Exception as e:
            error_type = type(e).__name__
            logger.error(f"Error extracting text from PDF {file_path} (type: {error_type}): {e}")
            
            # Fallback to simple extraction if structured extraction fails
            try:
                logger.info(f"Attempting fallback simple text extraction for {file_path}")
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                logger.info(f"Fallback extraction succeeded for {file_path}")
            except Exception as fallback_error:
                fallback_error_type = type(fallback_error).__name__
                error_msg = (
                    f"Failed to extract text from PDF '{file_path.name}'. "
                    f"Primary error ({error_type}): {str(e)}. "
                    f"Fallback error ({fallback_error_type}): {str(fallback_error)}. "
                    f"The PDF may be corrupted, encrypted, or in an unsupported format."
                )
                logger.error(error_msg)
                raise ValueError(error_msg) from e
        
        if not text_parts:
            error_msg = f"No text could be extracted from PDF '{file_path.name}'. The PDF may be empty, image-only, or corrupted."
            logger.warning(error_msg)
            # Return empty text instead of raising error - let the user know via metadata
            metadata["extraction_warning"] = error_msg
        
        return "\n\n".join(text_parts), metadata
    
    def _extract_text_with_structure(self, page) -> str:
        """Extract text from PDF page preserving structure (headings, lists, etc.).
        
        Args:
            page: pdfplumber Page object
            
        Returns:
            Structured text string
        """
        import pdfplumber
        
        try:
            # Extract words with their properties (font size, position)
            words = page.extract_words(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False,
                use_text_flow=True,
                extra_attrs=["size", "fontname"]
            )
        except Exception as e:
            logger.warning(f"Failed to extract words with structure, falling back to simple extraction: {e}")
            return page.extract_text() or ""
        
        if not words:
            # Fallback to simple text extraction
            return page.extract_text() or ""
        
        # Group words by line and detect structure
        lines = []
        current_line = []
        current_y = None
        current_font_size = None
        
        # Sort words by y position (top to bottom) and x position (left to right)
        # pdfplumber: top increases downward, so smaller top = higher on page
        words_sorted = sorted(words, key=lambda w: (w['top'], w['x0']))
        
        for word in words_sorted:
            word_y = word['top']
            font_size = word.get('size', 0)
            
            # Check if this word is on a new line (y position differs significantly)
            if current_y is None or abs(word_y - current_y) > 5:
                # Save previous line if exists
                if current_line:
                    line_text = ' '.join([w['text'] for w in current_line])
                    # Detect if this might be a heading (larger font or at top of page)
                    if current_font_size and current_font_size > 12:
                        lines.append(f"## {line_text}")
                    else:
                        lines.append(line_text)
                
                # Start new line
                current_line = [word]
                current_y = word_y
                current_font_size = font_size
            else:
                # Same line, add word
                current_line.append(word)
                # Update font size if this word has a larger font
                if font_size > (current_font_size or 0):
                    current_font_size = font_size
        
        # Add last line
        if current_line:
            line_text = ' '.join([w['text'] for w in current_line])
            if current_font_size and current_font_size > 12:
                lines.append(f"## {line_text}")
            else:
                lines.append(line_text)
        
        # Detect lists (lines starting with bullets, numbers, or dashes)
        structured_lines = []
        list_pattern = r'^[•\-\*\d]+[\.\)]\s+'
        for i, line in enumerate(lines):
            stripped = line.strip()
            # Check for list markers
            if re.match(list_pattern, stripped) or stripped.startswith('- ') or stripped.startswith('• '):
                # Remove list marker prefix
                cleaned_line = re.sub(list_pattern, '', stripped)
                structured_lines.append(f"  - {cleaned_line}")
            elif stripped.startswith('##'):
                structured_lines.append(stripped)
            else:
                structured_lines.append(stripped)
        
        return '\n'.join(structured_lines)
    
    def _format_table(self, table: List[List[Optional[str]]]) -> str:
        """Format a table extracted from PDF into readable text.
        
        Args:
            table: List of rows, each row is a list of cell values
            
        Returns:
            Formatted table as string
        """
        if not table:
            return ""
        
        # Filter out None values and clean cells
        cleaned_table = []
        for row in table:
            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
            # Only add row if it has at least one non-empty cell
            if any(cell for cell in cleaned_row):
                cleaned_table.append(cleaned_row)
        
        if not cleaned_table:
            return ""
        
        # Determine column widths
        num_cols = max(len(row) for row in cleaned_table) if cleaned_table else 0
        if num_cols == 0:
            return ""
        
        col_widths = [0] * num_cols
        for row in cleaned_table:
            for i, cell in enumerate(row[:num_cols]):
                col_widths[i] = max(col_widths[i], len(cell))
        
        # Format table
        formatted_rows = []
        for row_idx, row in enumerate(cleaned_table):
            # Pad row to match number of columns
            padded_row = row + [""] * (num_cols - len(row))
            
            # Format cells
            formatted_cells = []
            for i, cell in enumerate(padded_row[:num_cols]):
                formatted_cells.append(cell.ljust(col_widths[i]))
            
            formatted_rows.append(" | ".join(formatted_cells))
            
            # Add separator after header row (first row)
            if row_idx == 0 and len(cleaned_table) > 1:
                separator = "-" * (sum(col_widths) + 3 * (num_cols - 1))
                formatted_rows.append(separator)
        
        return "\n".join(formatted_rows)
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try with errors='replace'
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    
    def _extract_text_from_html(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Extract text
        text = soup.get_text(separator='\n', strip=True)
        
        # Clean up multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text
    
    def chunk_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """Chunk text into overlapping segments for embedding.
        
        For Excel files, chunks by sheet to preserve sheet context.
        For other documents, uses standard chunking with boundary detection.
        
        Args:
            text: Text to chunk
            metadata: Optional metadata to attach to each chunk (may include 'file_type' from document metadata)
            
        Returns:
            List of chunk dicts with keys: 'text', 'metadata', 'chunk_index'
        """
        if not text.strip():
            return []
        
        # Check if this is an Excel file by looking for sheet markers
        # Excel extraction format: "[Лист: SheetName]" followed by content
        file_type = metadata.get('file_type', '').lower() if metadata else ''
        is_excel = file_type in ['.xlsx', '.xls'] or '[Лист:' in text
        
        if is_excel:
            return self._chunk_excel_text(text, metadata)
        else:
            return self._chunk_standard_text(text, metadata)
    
    def _chunk_excel_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """Chunk Excel file text by sheet boundaries.
        
        Each sheet is chunked separately, and sheet name is preserved in metadata
        and included at the start of each chunk for better searchability.
        
        Args:
            text: Extracted Excel text with [Лист: SheetName] markers
            metadata: Optional metadata
            
        Returns:
            List of chunks with sheet context preserved
        """
        import re
        
        # Split text by sheet markers: [Лист: SheetName]
        sheet_pattern = r'\[Лист:\s*([^\]]+)\]'
        
        # Find all sheet markers and their positions
        sheet_matches = list(re.finditer(sheet_pattern, text))
        
        if not sheet_matches:
            # No sheet markers found, fall back to standard chunking
            logger.warning("Excel file has no sheet markers, using standard chunking")
            return self._chunk_standard_text(text, metadata)
        
        all_chunks = []
        chunk_index = 0
        
        # Process each sheet separately
        for i, match in enumerate(sheet_matches):
            sheet_name = match.group(1).strip()
            sheet_start = match.end()
            
            # Find where this sheet ends (start of next sheet or end of text)
            if i + 1 < len(sheet_matches):
                sheet_end = sheet_matches[i + 1].start()
            else:
                sheet_end = len(text)
            
            # Extract sheet content
            sheet_content = text[sheet_start:sheet_end].strip()
            
            if not sheet_content or sheet_content == '(пустой лист)':
                # Skip empty sheets but still create a minimal chunk for searchability
                chunk_metadata = {
                    "chunk_index": chunk_index,
                    "sheet_name": sheet_name,
                    "chunk_start": sheet_start,
                    "chunk_end": sheet_end,
                    **(metadata or {}),
                }
                all_chunks.append({
                    "text": f"[Лист: {sheet_name}]\n(пустой лист)",
                    "metadata": chunk_metadata,
                    "chunk_index": chunk_index,
                })
                chunk_index += 1
                continue
            
            # Chunk this sheet's content
            # Ensure sheet name is included in each chunk for better searchability
            sheet_chunks = self._chunk_single_sheet(sheet_content, sheet_name, chunk_index, sheet_start, metadata)
            
            all_chunks.extend(sheet_chunks)
            chunk_index += len(sheet_chunks)
        
        logger.info(f"Chunked Excel text into {len(all_chunks)} chunks across {len(sheet_matches)} sheets")
        return all_chunks
    
    def _chunk_single_sheet(self, sheet_content: str, sheet_name: str, start_chunk_index: int, 
                           sheet_start_pos: int, metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """Chunk a single sheet's content with sheet name preserved.
        
        Args:
            sheet_content: Content of one sheet (without [Лист:] header)
            sheet_name: Name of the sheet
            start_chunk_index: Starting chunk index
            sheet_start_pos: Character position where sheet starts in full text
            metadata: Base metadata
            
        Returns:
            List of chunks for this sheet
        """
        chunk_size_chars = self.chunk_size_tokens * 4
        chunk_overlap_chars = self.chunk_overlap_tokens * 4
        
        chunks = []
        start = 0
        chunk_index = start_chunk_index
        
        # Ensure sheet name header is included in the content
        sheet_header = f"[Лист: {sheet_name}]"
        
        while start < len(sheet_content):
            # Calculate end position
            end = start + chunk_size_chars - len(sheet_header) - 10  # Reserve space for header
            
            # Extract chunk content
            chunk_content = sheet_content[start:end] if end < len(sheet_content) else sheet_content[start:]
            
            # Try to break at paragraph boundary (table row boundaries are usually double newlines)
            if end < len(sheet_content):
                overlap_region = sheet_content[max(start, end - chunk_overlap_chars):end]
                para_break = overlap_region.rfind('\n\n')
                
                if para_break != -1:
                    end = max(start, end - chunk_overlap_chars) + para_break + 2
                    chunk_content = sheet_content[start:end]
            
            # Prepend sheet header to each chunk for better searchability
            full_chunk_text = f"{sheet_header}\n\n{chunk_content.strip()}"
            
            # Create chunk metadata with sheet name
            chunk_metadata = {
                "chunk_index": chunk_index,
                "sheet_name": sheet_name,
                "chunk_start": sheet_start_pos + start,
                "chunk_end": sheet_start_pos + min(end, len(sheet_content)),
                **(metadata or {}),
            }
            
            chunks.append({
                "text": full_chunk_text,
                "metadata": chunk_metadata,
                "chunk_index": chunk_index,
            })
            
            # Move start position (with overlap)
            if end >= len(sheet_content):
                break
            
            start = end - chunk_overlap_chars
            chunk_index += 1
        
        return chunks
    
    def _chunk_standard_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """Standard text chunking with boundary detection.
        
        Args:
            text: Text to chunk
            metadata: Optional metadata
            
        Returns:
            List of chunks
        """
        chunk_size_chars = self.chunk_size_tokens * 4
        chunk_overlap_chars = self.chunk_overlap_tokens * 4
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            # Calculate end position
            end = start + chunk_size_chars
            
            # Extract chunk
            chunk_text = text[start:end]
            
            # Try to break at sentence/paragraph boundary if possible
            if end < len(text):
                # Look for paragraph break (double newline) within overlap region
                overlap_region = text[max(start, end - chunk_overlap_chars):end]
                para_break = overlap_region.rfind('\n\n')
                
                if para_break != -1:
                    # Adjust end to paragraph boundary
                    end = max(start, end - chunk_overlap_chars) + para_break + 2
                    chunk_text = text[start:end]
                else:
                    # Look for sentence boundary (period + space/newline)
                    sentence_break = overlap_region.rfind('. ')
                    if sentence_break == -1:
                        sentence_break = overlap_region.rfind('.\n')
                    
                    if sentence_break != -1:
                        end = max(start, end - chunk_overlap_chars) + sentence_break + 2
                        chunk_text = text[start:end]
            
            # Create chunk metadata
            chunk_metadata = {
                "chunk_index": chunk_index,
                "chunk_start": start,
                "chunk_end": min(end, len(text)),
                **(metadata or {}),
            }
            
            chunks.append({
                "text": chunk_text.strip(),
                "metadata": chunk_metadata,
                "chunk_index": chunk_index,
            })
            
            # Move start position (with overlap)
            if end >= len(text):
                break
            
            start = end - chunk_overlap_chars
            chunk_index += 1
        
        logger.info(f"Chunked text into {len(chunks)} chunks (target size: {self.chunk_size_tokens} tokens, overlap: {self.chunk_overlap_tokens} tokens)")
        return chunks
    
    def get_text_preview(self, text: str, max_length: int = 10000) -> str:
        """Get a preview of text (truncated for display).
        
        Args:
            text: Full text
            max_length: Maximum length for preview
            
        Returns:
            Truncated text with ellipsis if needed
        """
        if len(text) <= max_length:
            return text
        
        # Try to truncate at sentence boundary
        truncated = text[:max_length]
        last_period = truncated.rfind('.')
        last_newline = truncated.rfind('\n')
        
        # Use the later boundary
        boundary = max(last_period, last_newline)
        
        if boundary > max_length * 0.8:  # Only use boundary if it's not too early
            return text[:boundary + 1] + "\n\n..."
        else:
            return text[:max_length] + "..."
    
    def _extract_text_with_vision_api(self, file_path: Path, page_num: Optional[int] = None) -> str:
        """Extract text from PDF page or image using GPT-4o Vision API via OpenRouter.
        
        Args:
            file_path: Path to PDF file or image
            page_num: Page number (for PDFs). If None, treats as single image.
            
        Returns:
            Extracted text string
        """
        try:
            from pdf2image import convert_from_path
            from PIL import Image
            import base64
            from io import BytesIO
        except ImportError:
            raise ImportError(
                "Vision API dependencies not installed. "
                "Install with: pip install pdf2image Pillow"
            )
        
        # Convert to image if PDF
        if file_path.suffix.lower() == ".pdf":
            if page_num is None:
                # Convert all pages
                images = convert_from_path(str(file_path), dpi=300)
                text_parts = []
                for img in images:
                    text = self._ocr_image_with_vision_api(img)
                    if text.strip():
                        text_parts.append(text)
                return "\n\n".join(text_parts)
            else:
                # Convert specific page
                images = convert_from_path(str(file_path), first_page=page_num, last_page=page_num, dpi=300)
                if images:
                    return self._ocr_image_with_vision_api(images[0])
                return ""
        else:
            # Image file
            img = Image.open(file_path)
            return self._ocr_image_with_vision_api(img)
    
    def _ocr_image_with_vision_api(self, image) -> str:
        """Extract text from PIL Image using GPT-4o Vision API.
        
        Args:
            image: PIL Image object
            
        Returns:
            Extracted text string
        """
        try:
            import base64
            from io import BytesIO
            from PIL import Image as PILImage
            from app.services.llm.client import LLMClient
            from app.core.database import SessionLocal
        except ImportError:
            raise ImportError("Required dependencies not available")
        
        # Convert image to base64
        buffered = BytesIO()
        image.save(buffered, format="PNG")
        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
        
        # Create LLM client
        db = SessionLocal()
        try:
            client = LLMClient(db=db)
            
            # Prepare vision message
            system_prompt = (
                "You are an OCR (Optical Character Recognition) assistant. "
                "Extract all text from the provided image. "
                "Preserve the structure, formatting, and layout as much as possible. "
                "If the image contains tables, format them clearly. "
                "If the text is in Russian, preserve Russian characters correctly. "
                "Return only the extracted text, without any additional commentary."
            )
            
            user_prompt = [
                {
                    "type": "text",
                    "text": "Extract all text from this image. Preserve structure and formatting."
                },
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{img_base64}"
                    }
                }
            ]
            
            # Get vision model from config
            from app.core.config import DEFAULT_VISION_MODEL
            vision_model = DEFAULT_VISION_MODEL
            
            # Call vision API
            response = client.client.chat.completions.create(
                model=vision_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.0,  # Deterministic output for OCR
                max_tokens=4000,  # Enough for most pages
            )
            
            extracted_text = response.choices[0].message.content
            return extracted_text if extracted_text else ""
            
        except Exception as e:
            logger.error(f"Vision API OCR failed: {e}")
            raise ValueError(f"Vision API OCR failed: {str(e)}") from e
        finally:
            db.close()
    
    def _extract_text_from_image(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from image file using AI OCR via OpenRouter.
        
        Args:
            file_path: Path to image file
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        metadata = {
            "filename": file_path.name,
            "file_size": file_path.stat().st_size,
            "file_type": file_path.suffix.lower(),
            "ocr_method": "ai",  # Always AI via OpenRouter
            "ocr_used": True,
        }
        
        try:
            text = self._extract_text_with_vision_api(file_path)
            
            metadata["text_length"] = len(text)
            metadata["extraction_success"] = True
            
            return text, metadata
        except Exception as e:
            logger.error(f"Failed to extract text from image {file_path}: {e}")
            metadata["extraction_success"] = False
            metadata["error"] = str(e)
            raise
    
    def _extract_text_from_excel(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Excel file (XLSX/XLS), converting all sheets to formatted text tables.
        
        Each sheet is extracted as a formatted table with headers and rows.
        This preserves structure for semantic search queries while making data queryable.
        
        Args:
            file_path: Path to Excel file (.xlsx or .xls)
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        import pandas as pd
        
        metadata = {
            "filename": file_path.name,
            "file_size": file_path.stat().st_size,
            "file_type": file_path.suffix.lower(),
            "sheet_count": 0,
            "total_rows": 0,
            "sheets": [],
        }
        
        text_parts = []
        
        try:
            # Read all sheets from Excel file
            # pd.read_excel() with sheet_name=None returns dict of {sheet_name: DataFrame}
            excel_file = pd.ExcelFile(file_path, engine='openpyxl' if file_path.suffix.lower() == '.xlsx' else None)
            all_sheets = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl' if file_path.suffix.lower() == '.xlsx' else None)
            
            metadata["sheet_count"] = len(all_sheets)
            
            # Process each sheet
            for sheet_name, df in all_sheets.items():
                sheet_metadata = {
                    "sheet_name": sheet_name,
                    "rows": len(df),
                    "columns": len(df.columns),
                }
                metadata["sheets"].append(sheet_metadata)
                metadata["total_rows"] += len(df)
                
                # Skip empty sheets
                if df.empty:
                    text_parts.append(f"[Лист: {sheet_name}]\n(пустой лист)\n")
                    continue
                
                # Convert DataFrame to formatted text table
                # Replace NaN with empty strings for cleaner output
                df_cleaned = df.fillna("")
                
                # Format as table (similar to PDF table formatting)
                # Use pipe-separated format for readability
                formatted_table = self._format_dataframe_table(df_cleaned, sheet_name)
                
                text_parts.append(formatted_table)
            
            combined_text = "\n\n" + "=" * 80 + "\n\n".join(text_parts)
            metadata["text_length"] = len(combined_text)
            metadata["extraction_success"] = True
            
            return combined_text, metadata
            
        except ImportError as e:
            error_msg = "pandas and openpyxl are required for Excel file processing. Install with: pip install pandas openpyxl"
            logger.error(f"{error_msg}: {e}")
            metadata["extraction_success"] = False
            metadata["error"] = error_msg
            raise ImportError(error_msg) from e
        except Exception as e:
            logger.error(f"Failed to extract text from Excel file {file_path}: {e}")
            metadata["extraction_success"] = False
            metadata["error"] = str(e)
            raise
    
    def _format_dataframe_table(self, df, sheet_name: str) -> str:
        """Format pandas DataFrame as a readable text table.
        
        Args:
            df: pandas DataFrame to format
            sheet_name: Name of the Excel sheet
            
        Returns:
            Formatted table string with headers and rows
        """
        import pandas as pd
        
        if df.empty:
            return f"[Лист: {sheet_name}]\n(пустой лист)"
        
        # Start with sheet name header
        lines = [f"[Лист: {sheet_name}]", ""]
        
        # Format column names as header row
        column_names = [str(col) for col in df.columns]
        
        # Convert all data to strings, handling different data types
        # Format dates, numbers, and preserve text
        formatted_rows = []
        for idx, row in df.iterrows():
            formatted_row = []
            for val in row:
                if pd.isna(val):
                    formatted_val = ""
                elif isinstance(val, (int, float)):
                    # Format numbers (remove unnecessary decimals for integers)
                    if isinstance(val, float) and val.is_integer():
                        formatted_val = str(int(val))
                    else:
                        formatted_val = str(val)
                elif isinstance(val, pd.Timestamp):
                    # Format dates consistently
                    formatted_val = val.strftime("%Y-%m-%d %H:%M:%S")
                else:
                    formatted_val = str(val).strip()
                
                formatted_row.append(formatted_val)
            
            formatted_rows.append(formatted_row)
        
        # Calculate column widths (including header)
        all_rows = [column_names] + formatted_rows
        num_cols = len(column_names)
        col_widths = [0] * num_cols
        
        for row in all_rows:
            for i, cell in enumerate(row[:num_cols]):
                col_widths[i] = max(col_widths[i], len(str(cell)))
        
        # Format header row
        header_cells = [str(col).ljust(col_widths[i]) for i, col in enumerate(column_names)]
        lines.append(" | ".join(header_cells))
        lines.append("-" * (sum(col_widths) + 3 * (num_cols - 1)))
        
        # Format data rows (limit to first 1000 rows to avoid huge documents)
        max_rows = 1000
        for row_idx, row in enumerate(formatted_rows[:max_rows]):
            row_cells = [str(cell).ljust(col_widths[i]) for i, cell in enumerate(row[:num_cols])]
            lines.append(" | ".join(row_cells))
        
        if len(formatted_rows) > max_rows:
            lines.append(f"\n... (показано {max_rows} из {len(formatted_rows)} строк)")
        
        return "\n".join(lines)

